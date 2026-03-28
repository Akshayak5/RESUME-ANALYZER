"""
Resume Parser: Extracts raw text from PDF and DOCX files.
"""
import os

def parse_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except ImportError:
        raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {str(e)}")


def parse_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {str(e)}")


def parse_txt(filepath: str) -> str:
    """Extract text from a plain text file."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def parse_resume(filepath: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(filepath)[1].lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,
        ".txt": parse_txt,
    }
    if ext not in parsers:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT")
    return parsers[ext](filepath)


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def is_allowed_file(filename: str) -> bool:
    ext = os.path.splitext(filename)[1].lower()
    return ext in ALLOWED_EXTENSIONS
