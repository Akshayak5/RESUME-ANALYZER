from flask import Blueprint, request, jsonify, current_app
from bson import ObjectId
from bson.errors import InvalidId
import traceback
import re
import json

resume_bp = Blueprint("resume", __name__)


def get_db():
    if hasattr(current_app._get_current_object(), 'mongo'):
        return current_app._get_current_object().mongo.db
    from backend.app import mongo
    return mongo.db


def safe_id(id_str):
    try:
        return ObjectId(str(id_str))
    except Exception:
        return None


def serialize_doc(doc):
    if doc is None:
        return None
    doc = dict(doc)
    if '_id' in doc:
        doc['_id'] = str(doc['_id'])
    if 'uploaded_at' in doc and hasattr(doc['uploaded_at'], 'isoformat'):
        doc['uploaded_at'] = doc['uploaded_at'].isoformat()
    if 'created_at' in doc and hasattr(doc['created_at'], 'isoformat'):
        doc['created_at'] = doc['created_at'].isoformat()
    if 'analysis' in doc and isinstance(doc['analysis'], dict):
        raw = doc['analysis'].get('score', 0)
        try:
            doc['analysis']['score'] = int(raw)
        except (TypeError, ValueError):
            doc['analysis']['score'] = 0
    return doc


def get_groq_client():
    import os
    from groq import Groq
    api_key = os.getenv("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("GROQ_API_KEY is not set in environment")
    return Groq(api_key=api_key)


def safe_parse_json(text):
    """
    Robustly parse JSON from LLM output.
    Handles markdown fences, control characters, and unescaped newlines.
    """
    # Strip markdown fences
    text = text.strip()
    if text.startswith("```"):
        parts = text.split("```")
        text = parts[1] if len(parts) > 1 else text
        if text.startswith("json"):
            text = text[4:]
    text = text.strip()

    # Remove null bytes and other problematic control characters
    # Keep \n (0x0a) and \r (0x0d) as they are valid in JSON outside strings
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # Try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try replacing literal newlines/tabs inside the JSON
    # This handles cases where the LLM puts real newlines inside string values
    cleaned = text.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Last resort: extract JSON object with regex
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            extracted = match.group(0).replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')
            return json.loads(extracted)

    raise ValueError("Could not parse JSON from LLM response")


# ══════════════════════════════════════════════════════════════
#  STATIC ROUTES — must come BEFORE dynamic /<id> routes
# ══════════════════════════════════════════════════════════════

@resume_bp.route("/analyze", methods=["POST"])
def analyze():
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file    = request.files["file"]
        user_id = request.form.get("user_id", "anonymous")
        ext     = file.filename.rsplit(".", 1)[-1].lower() if file.filename else ""

        if ext not in {"pdf", "docx", "doc", "txt"}:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 415

        from backend.utils.resume_parser import parse_resume
        from backend.utils.nlp_engine import (
            extract_skills, compute_score, get_missing_skills,
            recommend_jobs, extract_experience_years,
        )
        from backend.models.schemas import resume_schema

        file_bytes = file.read()
        parsed     = parse_resume(file_bytes, file.filename)

        raw_text_val = parsed.get("raw_text") or parsed.get("text") or ""
        if raw_text_val:
            parsed["raw_text"] = raw_text_val

        sd  = extract_skills(parsed.get("raw_text", ""))
        exp = extract_experience_years(parsed.get("raw_text", ""))

        analysis = {
            "skills":           sd.get("skills", []),
            "categorized":      sd.get("categorized", {}),
            "score":            int(compute_score(sd.get("skills", []), exp)),
            "missing_skills":   get_missing_skills(sd.get("skills", [])),
            "recommended_jobs": recommend_jobs(sd.get("skills", [])),
            "experience_years": int(exp),
            "verified_skills":  [],
        }

        doc = resume_schema(user_id, file.filename, parsed, analysis)
        db  = get_db()
        rid = str(db.resumes.insert_one(doc).inserted_id)

        return jsonify({
            "resume_id": rid,
            "candidate": parsed,
            "analysis":  analysis,
        }), 201

    except Exception as e:
        current_app.logger.error(f"analyze error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/history", methods=["GET"])
def history():
    try:
        user_id = request.args.get("user_id", "")
        db      = get_db()
        docs = list(
            db.resumes.find(
                {"user_id": user_id},
                {
                    "_id": 1, "filename": 1, "uploaded_at": 1,
                    "candidate": 1,
                    "analysis.score": 1,
                    "analysis.skills": 1,
                    "analysis.experience_years": 1,
                },
            )
            .sort("uploaded_at", -1)
            .limit(50)
        )
        return jsonify({"history": [serialize_doc(d) for d in docs]}), 200

    except Exception as e:
        current_app.logger.error(f"history error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/verify-skill", methods=["POST"])
def verify_skill():
    try:
        d      = request.get_json() or {}
        score  = int(d.get("score", 0))
        total  = int(d.get("total", 1))
        passed = (score / total) >= 0.6
        db     = get_db()

        from backend.models.schemas import test_result_schema
        db.test_results.insert_one(
            test_result_schema(d.get("user_id"), d.get("skill"), score, total)
        )

        if passed and d.get("resume_id"):
            oid = safe_id(d["resume_id"])
            if oid:
                db.resumes.update_one(
                    {"_id": oid},
                    {"$addToSet": {"analysis.verified_skills": d["skill"]}},
                )

        return jsonify({"passed": passed, "skill": d.get("skill")}), 200

    except Exception as e:
        current_app.logger.error(f"verify-skill error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/generate-questions", methods=["POST"])
def generate_questions():
    """
    Bug 3 fix: Generate skill-appropriate questions using Groq.
    Questions are tailored to the exact skill domain — not generic tech questions.
    Bug 2 fix: Rejects non-verifiable skills (soft skills / certifications).
    """
    try:
        from backend.utils.nlp_engine import get_verifiable_skills

        d     = request.get_json() or {}
        skill = (d.get("skill") or "").strip()

        if not skill:
            return jsonify({"error": "skill is required"}), 400

        # Bug 2: Reject unverifiable skills before generating questions
        verifiable = get_verifiable_skills([skill])
        if not verifiable:
            return jsonify({
                "error": f"'{skill}' is a soft skill or certification and cannot be quiz-tested.",
                "verifiable": False,
            }), 400

        client = get_groq_client()

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=2000,
            temperature=0.5,
            messages=[{
                "role": "system",
                "content": (
                    "You are a technical quiz generator. "
                    "Return ONLY valid JSON — no markdown, no explanation, no extra text."
                )
            }, {
                "role": "user",
                "content": f"""Generate exactly 5 multiple-choice quiz questions to test knowledge of: "{skill}".

CRITICAL RULES:
- Every question MUST be directly and specifically about "{skill}" — not about unrelated tech concepts.
- If "{skill}" is a design tool (e.g. Figma, Adobe XD): ask about its features, shortcuts, panels, components, prototyping.
- If "{skill}" is a programming language (e.g. Python, Java): ask about syntax, built-in functions, data types, error handling.
- If "{skill}" is a soft skill or domain knowledge (e.g. UX Design, Problem Solving): ask about processes, methodologies, best practices, real-world application of that skill.
- If "{skill}" is a healthcare/law/business skill: ask domain-specific scenario questions.
- NEVER ask generic questions like "what does API stand for" unless the skill is specifically about APIs.
- Each question must have exactly 4 answer options labeled A, B, C, D.
- Exactly one option must be correct.
- Difficulty should be appropriate for a professional resume verification (not too easy, not too hard).

Return ONLY this JSON:
{{
  "skill": "{skill}",
  "questions": [
    {{
      "question": "Question text here?",
      "options": ["A. Option one", "B. Option two", "C. Option three", "D. Option four"],
      "correct": "A"
    }}
  ]
}}"""
            }]
        )

        raw = response.choices[0].message.content
        result = safe_parse_json(raw)

        # Validate structure
        questions = result.get("questions", [])
        if not questions or len(questions) < 3:
            return jsonify({"error": "Failed to generate valid questions. Please try again."}), 500

        return jsonify({"skill": skill, "questions": questions}), 200

    except Exception as e:
        current_app.logger.error(f"generate-questions error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/verifiable-skills/<resume_id>", methods=["GET"])
def verifiable_skills(resume_id):
    """
    Bug 2 fix: Returns only the skills from a resume that can be meaningfully quiz-tested.
    The frontend should call this instead of showing ALL skills in 'Verify Your Skills'.
    """
    try:
        from backend.utils.nlp_engine import get_verifiable_skills

        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid ID: {resume_id}"}), 400

        db  = get_db()
        doc = db.resumes.find_one({"_id": oid}, {"analysis.skills": 1, "analysis.verified_skills": 1})
        if not doc:
            return jsonify({"error": "Resume not found"}), 404

        analysis         = doc.get("analysis") or {}
        all_skills       = analysis.get("skills", [])
        verified_skills  = analysis.get("verified_skills", [])
        quiz_skills      = get_verifiable_skills(all_skills)

        return jsonify({
            "verifiable_skills": quiz_skills,
            "verified_skills":   verified_skills,
        }), 200

    except Exception as e:
        current_app.logger.error(f"verifiable-skills error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/skills/trending", methods=["GET"])
def trending():
    try:
        db  = get_db()
        pipe = [
            {"$unwind": "$analysis.skills"},
            {"$group":  {"_id": "$analysis.skills", "count": {"$sum": 1}}},
            {"$sort":   {"count": -1}},
            {"$limit":  15},
        ]
        res = list(db.resumes.aggregate(pipe))
        return jsonify({
            "trending": [{"skill": r["_id"], "count": r["count"]} for r in res]
        }), 200

    except Exception as e:
        current_app.logger.error(f"trending error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════
#  DYNAMIC ROUTES — /<id> must be LAST
# ══════════════════════════════════════════════════════════════

@resume_bp.route("/<resume_id>", methods=["GET"])
def get_resume(resume_id):
    try:
        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid resume ID: {resume_id}"}), 400

        db  = get_db()
        doc = db.resumes.find_one({"_id": oid})
        if not doc:
            return jsonify({"error": "Resume not found"}), 404

        return jsonify(serialize_doc(doc)), 200

    except Exception as e:
        current_app.logger.error(f"get_resume error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/<resume_id>", methods=["PATCH"])
def update_resume(resume_id):
    try:
        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid ID: {resume_id}"}), 400

        data   = request.get_json() or {}
        db     = get_db()
        update = {}

        if data.get("candidate_name"):
            update["candidate.name"] = data["candidate_name"]
        if data.get("experience_years") is not None:
            update["analysis.experience_years"] = int(data["experience_years"])

        additional = data.get("additional_skills", [])
        if additional:
            result = db.resumes.update_one(
                {"_id": oid},
                {"$set": update, "$addToSet": {"analysis.skills": {"$each": additional}}},
            )
        else:
            result = db.resumes.update_one({"_id": oid}, {"$set": update})

        if result.matched_count == 0:
            return jsonify({"error": "Resume not found"}), 404

        return jsonify({"success": True, "modified": result.modified_count}), 200

    except Exception as e:
        current_app.logger.error(f"update_resume error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@resume_bp.route("/<resume_id>", methods=["DELETE"])
def delete_resume(resume_id):
    try:
        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid ID: {resume_id}"}), 400

        db     = get_db()
        result = db.resumes.delete_one({"_id": oid})

        if result.deleted_count == 0:
            return jsonify({"error": "Resume not found"}), 404

        return jsonify({"success": True, "deleted_id": resume_id}), 200

    except Exception as e:
        current_app.logger.error(f"delete_resume error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ── AI ENHANCE (Groq - free) ──────────────────────────────────────────────────
@resume_bp.route("/enhance/<resume_id>", methods=["POST"])
def enhance_resume(resume_id):
    try:
        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid ID: {resume_id}"}), 400

        db  = get_db()
        doc = db.resumes.find_one({"_id": oid})
        if not doc:
            return jsonify({"error": "Resume not found"}), 404

        candidate = doc.get("candidate") or {}
        raw_text = (
            candidate.get("raw_text")
            or doc.get("raw_text")
            or candidate.get("text")
            or doc.get("text")
            or ""
        ).strip()

        if not raw_text or len(raw_text) < 20:
            parts = [str(candidate.get(k)) for k in ["name", "email", "phone", "summary"] if candidate.get(k)]
            raw_text = " ".join(parts) if parts else json.dumps(candidate, default=str)[:3000]

        if not raw_text or len(raw_text) < 10:
            return jsonify({"error": "No resume text found. Please re-upload your resume."}), 400

        text_input = raw_text[:6000]
        client     = get_groq_client()

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=8000,
            temperature=0.3,
            messages=[{
                "role": "system",
                "content": "You are a professional resume editor. You MUST return ONLY valid JSON with no markdown, no explanation, no extra text. Do not include any control characters inside JSON string values."
            }, {
                "role": "user",
                "content": f"""You are a professional resume writer. Analyze this resume and return ONLY a valid JSON object.

Resume:
{text_input}

Rules:
- Read the resume carefully and extract EVERY section that exists in it — do not skip any section.
- Do NOT add sections that are not in the resume. Only include what is actually there.
- Professional Summary / Objective: Write exactly 3 sentences. Sentence 1: Who you are + your field/degree + years of experience or current status. Sentence 2: Your top 2-3 technical strengths and what you bring to the role. Sentence 3: Your career goal and the value you offer to an employer. Keep it concise, confident, and third-person free — write in first person implied (no 'I'). Each sentence must end with a period. No filler words like 'passionate', 'dynamic', 'results-driven'. Use specific, factual language.
- Experience / Work Experience: Each job as title line, then ">>>" prefixed bullets for key achievements (max 2-3 bullets, each max 15 words ending with a period).
- Projects: Project name with technology in brackets as title line, then ">>>" prefixed bullets for 2-3 key achievements (max 12 words each, ending with period).
- Education: One clean line per entry: "INSTITUTION NAME, Degree/Standard, Year, Grade" — nothing more.
- Certifications: One line per cert: "Cert Name - Provider (Month Year)."
- Activities / Extracurricular: Activity name as title line, then ">>>" prefixed 1-2 short bullets of what you did/gained.
- Skills / Programming Languages / Tools / Technologies / Soft Skills / Interests / Languages / Hobbies: Just the item names, nothing else.
- Awards / Achievements: Short one-line bullet per item ending with period.
- Any other section found in the resume: Format appropriately with short precise bullets.
- Fix ALL spelling and grammar mistakes.
- Replace weak phrases with strong action verbs (Led, Developed, Implemented, Designed, Achieved, Built, Managed).
- Do NOT use markdown like ** or * anywhere. Plain text only.
- Section headings must exactly match what is in the resume.

Return this JSON structure — enhanced_sections must contain ALL sections from the resume in order:
{{
  "spelling_fixes": [{{"original": "teh", "corrected": "the", "context": "context here"}}],
  "rephrased_bullets": [{{"original": "old text", "improved": "improved text"}}],
  "weak_phrases": [{{"phrase": "responsible for", "suggestion": "Led", "reason": "Use action verbs"}}],
  "summary_suggestion": "3-4 polished sentences ending with periods.",
  "overall_tips": ["tip 1", "tip 2", "tip 3"],
  "enhanced_sections": [
    {{
      "heading": "EXACT SECTION NAME FROM RESUME",
      "items": ["item 1", "item 2", ">>>sub detail 1", ">>>sub detail 2"]
    }}
  ],
  "enhanced_text": "Full improved resume as plain text with proper newlines between sections"
}}

CRITICAL:
- enhanced_sections MUST include every section from the resume — Education, Projects, Experience, Skills, Certifications, Activities, Awards, Languages, Interests, Hobbies, and anything else present.
- Professional Summary items = single string of 3-4 sentences with periods.
- All bullet items end with a period.
- ">>>" prefix = indented detail line below the main bullet above it."""
            }]
        )

        raw_response = response.choices[0].message.content
        result = safe_parse_json(raw_response)

        db.resumes.update_one(
            {"_id": oid},
            {"$set": {
                "candidate.enhanced_text":    result.get("enhanced_text", ""),
                "candidate.enhanced_sections": result.get("enhanced_sections", []),
            }}
        )

        return jsonify({"success": True, "enhancements": result}), 200

    except Exception as e:
        current_app.logger.error(f"enhance error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


# ── CAREER ROADMAP (Groq - free) ──────────────────────────────────────────────
@resume_bp.route("/roadmap", methods=["POST"])
def career_roadmap():
    try:
        data      = request.get_json() or {}
        job_title = (data.get("job_title") or "").strip()
        resume_id = (data.get("resume_id") or "").strip()

        if not job_title:
            return jsonify({"error": "job_title is required"}), 400

        current_skills = []
        if resume_id:
            oid = safe_id(resume_id)
            if oid:
                doc = get_db().resumes.find_one({"_id": oid})
                if doc:
                    current_skills = (doc.get("analysis") or {}).get("skills", [])

        skills_str = ", ".join(current_skills) if current_skills else "none provided"
        client     = get_groq_client()

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=2500,
            temperature=0.4,
            messages=[{
                "role": "system",
                "content": "You are a senior tech career coach. You MUST return ONLY valid JSON with no markdown, no explanation, no extra text. Do not include control characters inside string values."
            }, {
                "role": "user",
                "content": f"""Create a career roadmap for the role: "{job_title}".
Current skills: {skills_str}

Return ONLY this JSON structure:
{{
  "job_title": "{job_title}",
  "match_score": 42,
  "overview": "2-3 sentence assessment",
  "skills_have": ["skill1"],
  "skills_missing": [
    {{"skill": "PyTorch", "priority": "critical", "why": "Core framework", "learn_at": "fast.ai", "time_weeks": 4}}
  ],
  "phases": [
    {{
      "phase": 1,
      "title": "Foundation",
      "duration": "4-6 weeks",
      "goal": "What you can do after this phase",
      "tasks": [
        {{"task": "Task description", "resource": "Resource name", "type": "course", "hours": 20}}
      ]
    }}
  ],
  "projects": [
    {{"title": "Project name", "description": "what to build", "skills_practiced": ["skill1"], "difficulty": "beginner"}}
  ],
  "certifications": [
    {{"name": "Cert name", "provider": "Provider", "relevance": "High", "cost": "$150", "time_weeks": 8}}
  ],
  "salary_range": "$90,000 - $140,000",
  "time_to_ready": "4-6 months",
  "job_boards": ["LinkedIn", "Indeed"],
  "interview_topics": ["topic1", "topic2", "topic3"]
}}"""
            }]
        )

        raw = response.choices[0].message.content
        result = safe_parse_json(raw)
        return jsonify({"success": True, "roadmap": result}), 200

    except Exception as e:
        current_app.logger.error(f"roadmap error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500

# ── DOWNLOAD ENHANCED RESUME AS DOCX ─────────────────────────────────────────
@resume_bp.route("/download-enhanced/<resume_id>", methods=["GET"])
def download_enhanced(resume_id):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io
        from flask import send_file

        oid = safe_id(resume_id)
        if not oid:
            return jsonify({"error": f"Invalid ID: {resume_id}"}), 400

        db  = get_db()
        doc = db.resumes.find_one({"_id": oid})
        if not doc:
            return jsonify({"error": "Resume not found"}), 404

        candidate     = doc.get("candidate") or {}
        sections_data = candidate.get("enhanced_sections", [])
        enhanced_text = candidate.get("enhanced_text", "").strip()
        raw_text      = candidate.get("raw_text", "").strip()

        if not sections_data and not enhanced_text and not raw_text:
            return jsonify({"error": "No enhanced resume found. Please run Enhance first."}), 400

        name    = candidate.get("name", "Resume")
        email   = candidate.get("email", "")
        phone   = candidate.get("phone", "")

        # ── Build DOCX ────────────────────────────────────────────────────
        docx_doc = DocxDocument()

        # Margins — match typical resume (0.75 inch sides)
        for sec in docx_doc.sections:
            sec.top_margin    = Inches(0.75)
            sec.bottom_margin = Inches(0.75)
            sec.left_margin   = Inches(0.85)
            sec.right_margin  = Inches(0.85)

        # ── Helper: section heading with bottom border ────────────────────
        def add_heading(text):
            p = docx_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x1f, 0x39, 0x64)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1F3964")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── Helper: bullet item ───────────────────────────────────────────
        def add_bullet(text):
            clean = text.lstrip("-•▸ ").strip()
            if clean.startswith(">>>"):
                # Plain indented text — no bullet, no bold
                p = docx_doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.left_indent  = Inches(0.25)
                run = p.add_run(clean[3:].strip().replace("**", ""))
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            else:
                # Normal bullet — no bold
                p = docx_doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(0)
                run = p.add_run(clean.replace("**", ""))
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.bold = False

        # ── Helper: plain paragraph ───────────────────────────────────────
        def add_plain(text, italic=False, bold=False, center=False):
            p = docx_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.italic = italic
            run.bold   = bold
            return p

        # ══ HEADER — Name + Contact ════════════════════════════════════════
        # Name
        name_para = docx_doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        nr = name_para.add_run(name.upper())
        nr.bold = True
        nr.font.size = Pt(18)
        nr.font.name = "Calibri"
        nr.font.color.rgb = RGBColor(0x1f, 0x39, 0x64)

        # Contact info line
        contact_bits = []
        if phone:  contact_bits.append(phone)
        if email:  contact_bits.append(email)
        raw_lines = [l.strip() for l in (enhanced_text or raw_text).split("\n") if l.strip()]
        # Try to grab location from raw text (usually near top)
        for ln in raw_lines[:5]:
            if any(w in ln.lower() for w in ["india", "kerala", "bangalore", "bengaluru", "mumbai", "delhi"]):
                contact_bits.insert(0, ln)
                break

        if contact_bits:
            cp = docx_doc.add_paragraph(" | ".join(contact_bits))
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(2)
            for r in cp.runs:
                r.font.size = Pt(10)
                r.font.name = "Calibri"
                r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        # LinkedIn line if found
        for ln in raw_lines[:8]:
            if "linkedin" in ln.lower():
                lp = docx_doc.add_paragraph(ln.strip())
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lp.paragraph_format.space_after = Pt(4)
                for r in lp.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0x0, 0x70, 0xC0)
                break

        # ══ SECTIONS ═══════════════════════════════════════════════════════
        # Sections where items should be a plain sentence (not bullets)
        sentence_sections = {
            "professional summary", "summary", "objective", "profile"
        }
        # Sections where items should be comma-joined on one line
        inline_sections = {
            "programming languages", "tools", "tools and technologies",
            "soft skills", "skills", "technical skills",
            "languages", "interests"
        }

        if sections_data:
            for sec in sections_data:
                heading = (sec.get("heading") or "").strip()
                items   = [i.strip() for i in (sec.get("items") or []) if i.strip()]
                if not heading and not items:
                    continue

                add_heading(heading)
                hl = heading.lower()

                if hl in sentence_sections or any(w in hl for w in ['summary', 'objective', 'profile']):
                    # Summary — plain paragraph sentence
                    add_plain(" ".join(items))

                elif any(kw in hl for kw in ["skill", "language", "tool", "technolog", "soft skill", "interest", "hobbies", "hobby"]):
                    # Skills-type sections — each item as its own bullet
                    for item in items:
                        add_bullet(item.replace("**", ""))

                else:
                    # All other sections — bullets with >>> as indented plain text
                    for item in items:
                        add_bullet(item.replace("**", ""))

        else:
            # Fallback: parse enhanced_text / raw_text
            text = enhanced_text or raw_text
            section_kw = {
                "professional summary", "summary", "education", "experience",
                "projects", "certifications", "activities", "achievements",
                "programming languages", "tools and technologies", "soft skills",
                "interests", "languages", "skills",
            }
            current_section = ""
            for line in [l.strip() for l in text.split("\n") if l.strip()]:
                lower = line.lower().rstrip(":")
                if lower in section_kw:
                    current_section = lower
                    add_heading(line)
                elif line.startswith("-") or line.startswith("•"):
                    if current_section in sentence_sections or any(w in current_section for w in ['summary', 'objective', 'profile']):
                        add_plain(line.lstrip("-• ").strip())
                    else:
                        add_bullet(line)
                else:
                    if current_section in inline_sections:
                        add_plain(line)
                    elif current_section in sentence_sections or any(w in current_section for w in ['summary', 'objective', 'profile']):
                        add_plain(line)
                    else:
                        add_plain(line)

        # ── Save ──────────────────────────────────────────────────────────
        buf = io.BytesIO()
        docx_doc.save(buf)
        buf.seek(0)

        filename = f"enhanced_resume_{name.replace(' ', '_')}.docx"
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        current_app.logger.error(f"download-enhanced error: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500